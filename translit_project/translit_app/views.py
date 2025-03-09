from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.core.files.storage import FileSystemStorage
from django.utils import translation
from .utils import convert_to_cyrillic, convert_to_latin
from .models import TranslationHistory, Advertisement
import docx
import os
import chardet
from io import BytesIO
import re
from django.conf import settings
from .models import ContactSubmission
from django.contrib import messages


def clean_text(text):
    return ''.join(c for c in text if c.isprintable() or c in '\n\r\t')

def transliterate_words(text, direction):
    words = re.findall(r'\w+', text)
    trans_func = convert_to_cyrillic if direction == "to_cyrillic" else convert_to_latin
    trans_dict = {word: trans_func(word) for word in words}
    for word, trans in trans_dict.items():
        text = re.sub(r'\b' + word + r'\b', trans, text)
    return text

def index(request):
    result = ""
    error = ""
    history = TranslationHistory.objects.all().order_by('-created_at')[:5]
    ads = Advertisement.objects.filter(active=True).order_by('-created_at')

    if "lang" in request.GET:
        lang = request.GET["lang"]
        translation.activate(lang)
        request.session['django_language'] = lang

    if request.method == "POST":
        direction = request.POST.get("direction", "to_cyrillic")
        text = ""
        file_extension = None
        doc_content = None

        if "file" in request.FILES:
            file = request.FILES["file"]
            allowed_extensions = ['.txt', '.docx']
            max_size = 5 * 1024 * 1024
            file_extension = os.path.splitext(file.name)[1].lower()

            if file_extension not in allowed_extensions:
                error = "Faqat .txt yoki .docx fayllarni yuklashingiz mumkin!"
            elif file.size > max_size:
                error = "Fayl hajmi 5 MB dan oshmasligi kerak!"
            else:
                try:
                    fs = FileSystemStorage()
                    cleaned_filename = re.sub(r'[^a-zA-Z0-9_\.]', '_', file.name)
                    filename = fs.save(cleaned_filename, file)
                    file_path = fs.path(filename)

                    if file_extension == '.txt':
                        with open(file_path, 'rb') as f:
                            raw_data = f.read()
                        detected = chardet.detect(raw_data)
                        encoding = detected['encoding'] if detected['confidence'] > 0.7 else 'latin-1'
                        try:
                            text = raw_data.decode(encoding)
                        except UnicodeDecodeError:
                            for fallback_encoding in ['utf-8', 'latin-1', 'cp1251']:
                                try:
                                    text = raw_data.decode(fallback_encoding)
                                    break
                                except UnicodeDecodeError:
                                    continue
                            else:
                                error = "Faylni o‘qib bo‘lmadi: Kodlash xatosi!"
                    elif file_extension == '.docx':
                        doc = docx.Document(file_path)
                        doc_content = doc
                        text = "\n".join([para.text for para in doc.paragraphs])

                    os.remove(file_path)
                except Exception as e:
                    error = f"Faylni o‘qishda xatolik: {str(e)}"
        else:
            text = request.POST.get("text", "")

        if text and not error:
            try:
                if file_extension == '.docx':
                    result = transliterate_words(text, direction)
                    if doc_content:
                        for para in doc_content.paragraphs:
                            para.text = transliterate_words(para.text, direction)
                        result_doc = doc_content
                else:
                    result = convert_to_cyrillic(text) if direction == "to_cyrillic" else convert_to_latin(text)
                    result_doc = None

                TranslationHistory.objects.create(
                    original_text=text,
                    translated_text=result,
                    direction=direction,
                    file_name=file.name if "file" in request.FILES else None,
                    file_extension=file_extension
                )
            except Exception as e:
                error = f"Transliteratsiyada xatolik: {str(e)}"

    for ad in ads:
        ad.views += 1
        ad.save()

    return render(request, "translit_app/index.html", {
        "result": result,
        "history": history,
        "error": error,
        "ads": ads
    })

def download_result(request):
    if request.method == "POST":
        result = request.POST.get("result", "")
        if not result:
            return HttpResponse("Yuklab olish uchun natija mavjud emas!", status=400)

        last_history = TranslationHistory.objects.order_by('-created_at').first()
        file_extension = last_history.file_extension if last_history and last_history.file_extension else '.txt'

        try:
            cleaned_result = clean_text(result)
            if file_extension == '.txt':
                response = HttpResponse(cleaned_result, content_type="text/plain; charset=utf-8")
                response["Content-Disposition"] = 'attachment; filename="translated_text.txt"'
            elif file_extension == '.docx':
                doc = docx.Document()
                doc.add_paragraph(cleaned_result)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                response = HttpResponse(buffer.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                response["Content-Disposition"] = 'attachment; filename="translated_text.docx"'
                buffer.close()
            else:
                response = HttpResponse("Faqat .txt va .docx formatlari qo‘llab-quvvatlanadi!", status=400)
            return response
        except Exception as e:
            return HttpResponse(f"Yuklab olishda xatolik: {str(e)}", status=500)
    return render(request, "translit_app/index.html")

def translit_api(request):
    if request.method == "POST":
        text = request.POST.get("text", "")
        direction = request.POST.get("direction", "to_cyrillic")
        if text:
            try:
                result = convert_to_cyrillic(text) if direction == "to_cyrillic" else convert_to_latin(text)
                return JsonResponse({"result": result})
            except Exception as e:
                return JsonResponse({"error": f"Transliteratsiyada xatolik: {str(e)}"}, status=500)
        return JsonResponse({"error": "Matn kiritilmadi!"}, status=400)
    return JsonResponse({"error": "Faqat POST so‘rovlari qabul qilinadi!"}, status=405)

def features(request):
    ads = Advertisement.objects.all()
    return render(request, 'translit_app/features.html', {'ads': ads})

def contact(request):
    ads = Advertisement.objects.all()
    if request.method == "POST":
        # Forma ma'lumotlarini olish
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        message = request.POST.get('message')
        file = request.FILES.get('file')  # Fayl ixtiyoriy, shuning uchun None bo'lishi mumkin

        # Ma'lumotlarni tekshirish (ixtiyoriy qo'shimcha validatsiya)
        if not all([first_name, last_name, email, phone, message]):
            messages.error(request, "Barcha majburiy maydonlarni to‘ldiring!")
            return redirect('contact')

        # Fayl hajmini tekshirish (agar fayl yuklansa)
        if file and file.size > 5 * 1024 * 1024:  # 5 MB dan katta bo‘lsa
            messages.error(request, "Fayl hajmi 5 MB dan oshmasligi kerak!")
            return redirect('contact')

        try:
            # Ma’lumotlarni saqlash
            contact_submission = ContactSubmission(
                first_name=first_name,
                last_name=last_name,
                email=email,
                phone=phone,
                message=message,
                file=file
            )
            contact_submission.save()

            messages.success(request, "Xabaringiz muvaffaqiyatli yuborildi!")
            return redirect('contact')

        except Exception as e:
            # Xatolik yuz bersa, foydalanuvchiga xabar ko‘rsatish
            messages.error(request, f"Xatolik yuz berdi: {str(e)}")
            return redirect('contact')

    # GET so‘rov bo‘lsa, faqat shablonni ko‘rsatish
    return render(request, 'translit_app/contact.html', {'ads': ads})

    